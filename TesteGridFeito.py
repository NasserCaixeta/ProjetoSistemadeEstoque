import mysql.connector
import tkinter as tk
from tkinter import ttk, messagebox, filedialog
from docx import Document

class Aplicacao:
    def __init__(self, root):
        self.root = root
        self.root.title("Grid do Banco de Dados")
        self.root.geometry("1000x500")  # Defina o tamanho da janela conforme necessário

        # Conectar ao banco de dados
        self.conexao = mysql.connector.connect(
            host="localhost",
            user="root",
            password="senha",
            database="banco_dados_estoque"
        )

        # Criar um cursor para executar consultas SQL
        self.cursor = self.conexao.cursor()

        # Criar a primeira Treeview para o estoque
        self.treeview_estoque = ttk.Treeview(self.root, columns=("Caixa", "Unidade", "Quilo", "Litro", "Grama", "Prioridade"))
        self.treeview_estoque.heading("#0", text="Item")
        self.treeview_estoque.heading("Caixa", text="Caixa")
        self.treeview_estoque.heading("Unidade", text="Unidade")
        self.treeview_estoque.heading("Quilo", text="Quilo")
        self.treeview_estoque.heading("Litro", text="Litro")
        self.treeview_estoque.heading("Grama", text="Grama")
        self.treeview_estoque.heading("Prioridade", text="Prioridade")
        self.treeview_estoque.grid(row=0, column=0, pady=10, padx=10, sticky=tk.NSEW)

        # Criar a segunda Treeview para a lista para saída
        self.treeview_saida = ttk.Treeview(self.root, columns=("Caixa", "Unidade", "Quilo", "Litro", "Grama", "Prioridade"))
        self.treeview_saida.heading("#0", text="Item")
        self.treeview_saida.heading("Caixa", text="Caixa")
        self.treeview_saida.heading("Unidade", text="Unidade")
        self.treeview_saida.heading("Quilo", text="Quilo")
        self.treeview_saida.heading("Litro", text="Litro")
        self.treeview_saida.heading("Grama", text="Grama")
        self.treeview_saida.heading("Prioridade", text="Prioridade")
        self.treeview_saida.grid(row=0, column=2, pady=10, padx=10, sticky=tk.NSEW)

        # Adicionar Combobox
        self.criar_combobox()

        # Adicionar elementos para exibir informações
        self.criar_elementos_exibicao()

        # Adicionar botão para enviar itens para o Treeview
        self.botao_enviar = tk.Button(self.root, text="Enviar para o Grid", command=self.enviar_para_grid)
        self.botao_enviar.grid(row=9, column=1, pady=10, padx=10, sticky=tk.E)

        # Adicionar evento de clique na Treeview do estoque
        self.treeview_estoque.bind("<ButtonRelease-1>", self.mostrar_dados_selecionados)

        # Botão para exportar para o Word
        self.botao_exportar_word = tk.Button(self.root, text="Exportar para Word", command=self.exportar_para_word)
        self.botao_exportar_word.grid(row=9, column=2, pady=10, padx=10, sticky=tk.E)

        # Ajustar as colunas para ocupar menos espaço
        for column in ("Caixa", "Unidade", "Quilo", "Litro", "Grama", "Prioridade"):
            self.treeview_estoque.column(column, width=30, anchor="center")
            self.treeview_saida.column(column, width=30, anchor="center")

        # Preencher a Treeview do estoque com os itens do banco de dados
        self.preencher_treeview_estoque()

    def criar_combobox(self):
        # Consultar todos os nomes dos itens
        self.cursor.execute("SELECT Nome FROM itens")
        nomes = [nome[0] for nome in self.cursor.fetchall()]

        # Criar a Combobox
        self.combobox = ttk.Combobox(self.root, values=nomes)
        self.combobox.set("Escolha um item")  # Texto padrão exibido na Combobox
        self.combobox.grid(row=1, column=1, pady=10, padx=10, sticky=tk.W)


        # Adicionar evento para a mudança na seleção
        self.combobox.bind("<<ComboboxSelected>>", self.atualizar_informacoes)

        # Adicionar evento para autocompletar enquanto digita
        self.combobox.bind("<KeyPress>", self.autocompletar_combobox)

    def autocompletar_combobox(self, event):
        # Obter o texto digitado na Combobox
        entrada = self.combobox.get().lower()

        # Consultar todos os nomes dos itens
        self.cursor.execute("SELECT Nome FROM itens")
        nomes = [nome[0] for nome in self.cursor.fetchall()]

        # Filtrar nomes que começam com o texto digitado
        sugestoes = [nome for nome in nomes if nome.lower().startswith(entrada)]

        # Atualizar a lista de sugestões da Combobox
        self.combobox['values'] = sugestoes

    def criar_elementos_exibicao(self):
        # Criar Labels e Entries para exibir informações
        self.label_caixa = tk.Label(self.root, text="Caixa:")
        self.entry_caixa = tk.Entry(self.root)
        self.label_unidade = tk.Label(self.root, text="Unidade:")
        self.entry_unidade = tk.Entry(self.root)
        self.label_quilo = tk.Label(self.root, text="Quilo:")
        self.entry_quilo = tk.Entry(self.root)
        self.label_litro = tk.Label(self.root, text="Litro:")
        self.entry_litro = tk.Entry(self.root)
        self.label_grama = tk.Label(self.root, text="Grama:")
        self.entry_grama = tk.Entry(self.root)
        self.label_prioridade = tk.Label(self.root, text="Prioridade:")
        self.entry_prioridade = tk.Entry(self.root)

        # Posicionar Labels e Entries
        self.label_caixa.grid(row=3, column=0, pady=5, padx=10, sticky=tk.E)
        self.entry_caixa.grid(row=3, column=1, pady=5, padx=10, sticky=tk.E)
        self.label_unidade.grid(row=4, column=0, pady=5, padx=10, sticky=tk.E)
        self.entry_unidade.grid(row=4, column=1, pady=5, padx=10, sticky=tk.E)
        self.label_quilo.grid(row=5, column=0, pady=5, padx=10, sticky=tk.E)
        self.entry_quilo.grid(row=5, column=1, pady=5, padx=10, sticky=tk.E)
        self.label_litro.grid(row=6, column=0, pady=5, padx=10, sticky=tk.E)
        self.entry_litro.grid(row=6, column=1, pady=5, padx=10, sticky=tk.E)
        self.label_grama.grid(row=7, column=0, pady=5, padx=10, sticky=tk.E)
        self.entry_grama.grid(row=7, column=1, pady=5, padx=10, sticky=tk.E)
        self.label_prioridade.grid(row=8, column=0, pady=5, padx=10, sticky=tk.E)
        self.entry_prioridade.grid(row=8, column=1, pady=5, padx=10, sticky=tk.E)


    def atualizar_informacoes(self, event):
        # Obter o nome selecionado na ComboBox
        nome_selecionado = self.combobox.get()

        # Consultar informações do item selecionado
        self.cursor.execute("SELECT Caixa, Unidade, Quilo, Litro, Grama, Prioridade FROM itens WHERE Nome=%s", (nome_selecionado,))
        informacoes = self.cursor.fetchone()

        # Preencher as Entries com as informações obtidas
        self.entry_caixa.delete(0, tk.END)
        self.entry_caixa.insert(0, informacoes[0])

        self.entry_unidade.delete(0, tk.END)
        self.entry_unidade.insert(0, informacoes[1])

        self.entry_quilo.delete(0, tk.END)
        self.entry_quilo.insert(0, informacoes[2])

        self.entry_litro.delete(0, tk.END)
        self.entry_litro.insert(0, informacoes[3])

        self.entry_grama.delete(0, tk.END)
        self.entry_grama.insert(0, informacoes[4])

        self.entry_prioridade.delete(0, tk.END)
        self.entry_prioridade.insert(0, informacoes[5])

    def enviar_para_grid(self):
        nome_selecionado = self.combobox.get()

        # Obter os valores alterados das Entries
        caixa = self.entry_caixa.get()
        unidade = self.entry_unidade.get()
        quilo = self.entry_quilo.get()
        litro = self.entry_litro.get()
        grama = self.entry_grama.get()
        prioridade = self.entry_prioridade.get()

        # Validar se a prioridade é um número válido e não zero
        if prioridade and prioridade.isdigit() and int(prioridade) != 0:
            # Preencher a Treeview da lista para saída com os novos valores e organizar por prioridade
            self.treeview_saida.insert("", "end", text=nome_selecionado, values=(caixa, unidade, quilo, litro, grama, prioridade))
            self.ordenar_por_prioridade(self.treeview_saida)
        else:
            messagebox.showwarning("Aviso", "A prioridade deve ser um número válido e diferente de zero.")

    def ordenar_por_prioridade(self, treeview):
        # Organizar o Treeview pela coluna "Prioridade"
        data = [(treeview.item(item)["values"], item) for item in treeview.get_children()]
        data.sort(key=lambda x: int(x[0][5]))  # Ordenar com base na prioridade (índice 5)
        for index, (values, item) in enumerate(data):
            treeview.move(item, "", index)

    def exportar_para_word(self):
        file_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Documentos Word", "*.docx")])

        if file_path:
            # Criar um documento Word
            document = Document()

            # Adicionar cabeçalho para a lista para saída
            document.add_heading('Lista para Saída', level=1)

            # Adicionar dados do Treeview da lista para saída ao documento Word
            self.adicionar_dados_ao_documento(document, self.treeview_saida)

            # Salvar o documento Word no local escolhido pelo usuário
            document.save(file_path)
            messagebox.showinfo("Exportação Concluída", f"Os dados foram exportados para o arquivo {file_path}.")

    def adicionar_dados_ao_documento(self, document, treeview):
        for item in treeview.get_children():
            nome = treeview.item(item, "text")
            values = treeview.item(item, "values")

            # Verificar se caixas, unidades, quilos, litros e gramas são diferentes de "0"
            if any(value != "0" for value in values[0:5]):
                document.add_paragraph(f'Item: {nome}')
                
                if values[0] != "0":
                    document.add_paragraph(f'Caixa: {values[0]}')
                if values[1] != "0":
                    document.add_paragraph(f'Unidade: {values[1]}')
                if values[2] != "0":
                    document.add_paragraph(f'Quilo: {values[2]}')
                if values[3] != "0":
                    document.add_paragraph(f'Litro: {values[3]}')
                if values[4] != "0":
                    document.add_paragraph(f'Grama: {values[4]}')
                
                document.add_paragraph('\n')


    def mostrar_dados_selecionados(self, event):
        # Obter o item selecionado na Treeview do estoque
        item_selecionado = self.treeview_estoque.selection()

        if item_selecionado:
            # Obter o nome do item selecionado
            nome_selecionado = self.treeview_estoque.item(item_selecionado, "text")

            # Definir o nome selecionado na Combobox
            self.combobox.set(nome_selecionado)

            # Consultar informações do item selecionado
            self.cursor.execute("SELECT Caixa, Unidade, Quilo, Litro, Grama, Prioridade FROM itens WHERE Nome=%s", (nome_selecionado,))
            informacoes = self.cursor.fetchone()

            # Preencher as Entries com as informações obtidas
            self.entry_caixa.delete(0, tk.END)
            self.entry_caixa.insert(0, informacoes[0])

            self.entry_unidade.delete(0, tk.END)
            self.entry_unidade.insert(0, informacoes[1])

            self.entry_quilo.delete(0, tk.END)
            self.entry_quilo.insert(0, informacoes[2])

            self.entry_litro.delete(0, tk.END)
            self.entry_litro.insert(0, informacoes[3])

            self.entry_grama.delete(0, tk.END)
            self.entry_grama.insert(0, informacoes[4])

            self.entry_prioridade.delete(0, tk.END)
            self.entry_prioridade.insert(0, informacoes[5])

    def preencher_treeview_estoque(self):
        # Limpar a Treeview do estoque
        for item in self.treeview_estoque.get_children():
            self.treeview_estoque.delete(item)

        # Consultar todos os itens do banco de dados
        self.cursor.execute("SELECT Nome, Caixa, Unidade, Quilo, Litro, Grama, Prioridade FROM itens")
        for row in self.cursor.fetchall():
            nome, caixa, unidade, quilo, litro, grama, prioridade = row

            # Verificar se a unidade é diferente de '0' antes de adicionar à Treeview
            if unidade != '0':
                self.treeview_estoque.insert("", "end", text=nome, values=(caixa, unidade, quilo, litro, grama, prioridade))

if __name__ == "__main__":
    root = tk.Tk()
    app = Aplicacao(root)
    root.mainloop()
