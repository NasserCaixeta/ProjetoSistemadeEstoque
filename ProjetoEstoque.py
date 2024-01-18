import tkinter as tk
from tkinter import ttk, messagebox, Entry, Label, filedialog, Button, Frame, Tk
import mysql.connector
import docx
from ttkthemes import ThemedTk
from PIL import Image, ImageTk
import os
theme_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), "Azure Tema", "azure.tcl")
import os
icon_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), "IconePrograma.ico")


class ConexaoBancoDados:
    def __init__(self):
        # Primeiro, criamos o banco de dados e as tabelas se não existirem
        if not self.criar_banco_dados():
            # Se não foi possível criar o banco de dados, exibir uma mensagem de erro e encerrar o programa
            messagebox.showerror("Erro", "Não foi possível criar o banco de dados.")
            exit()

        # Agora, podemos estabelecer a conexão
        self.conexao = mysql.connector.connect(
            host="localhost",
            user="root",
            password="senha",
            database="banco_dados_estoque"
        )

        # Cria as tabelas se não existirem
        self.criar_tabela_itens()
        self.criar_tabela_alimentos()

    def verificar_banco_dados(self):
        try:
            cursor = self.obter_cursor()
            cursor.execute("SELECT 1 FROM itens LIMIT 1")
            result = cursor.fetchone()
            cursor.close()
            return result is not None
        except mysql.connector.Error as e:
            return False

    def criar_banco_dados(self):
        try:
            # Conectar sem especificar um banco de dados
            conexao_temp = mysql.connector.connect(
                host="localhost",
                user="root",
                password="senha"
            )
            cursor = conexao_temp.cursor()

            # Criar o banco de dados se não existir
            cursor.execute("CREATE DATABASE IF NOT EXISTS banco_dados_estoque")

            # Fechar a conexão temporária
            cursor.close()
            conexao_temp.close()

            return True  # Retorna True se o banco de dados foi criado com sucesso
        except mysql.connector.Error as e:
            # Mostrar uma mensagem de erro
            messagebox.showerror("Erro", f"Não foi possível criar o banco de dados: {e}")
            return False  # Retorna False se houve um erro ao criar o banco de dados

    def criar_tabela_itens(self):
        try:
            cursor = self.obter_cursor()
            cursor.execute('''
                CREATE TABLE IF NOT EXISTS itens (
                    ID INT AUTO_INCREMENT PRIMARY KEY,
                    Nome VARCHAR(255) NOT NULL,
                    Caixa INT,
                    Unidade INT,
                    Quilo INT,
                    Litro INT,
                    Grama INT,
                    Prioridade INT
                )
            ''')
            cursor.close()
            self.conexao.commit()
        except mysql.connector.Error as e:
            messagebox.showerror("Erro", f"Não foi possível criar a tabela 'itens': {e}")

    def criar_tabela_alimentos(self):
        try:
            cursor = self.obter_cursor()
            cursor.execute('''
                CREATE TABLE IF NOT EXISTS alimentos (
                    ID INT AUTO_INCREMENT PRIMARY KEY,
                    Nome VARCHAR(255) NOT NULL,
                    Caixa INT,
                    Unidade INT,
                    Quilo INT,
                    Litro INT,
                    Grama INT,
                    Prioridade INT
                )
            ''')
            cursor.close()
            self.conexao.commit()
        except mysql.connector.Error as e:
            messagebox.showerror("Erro", f"Não foi possível criar a tabela 'alimentos': {e}")

    def obter_cursor(self):
        return self.conexao.cursor()
    
    def fechar_conexao(self):
        self.conexao.close()
    
    def commit(self):
        self.conexao.commit()

    def execute(self, query, values=None):
        cursor = self.obter_cursor()
        if values:
            cursor.execute(query, values)
        else:
            cursor.execute(query)
        return cursor
    


class TelaInsercaoItens(ttk.Frame):
    def __init__(self, master=None, conexao=None):
        super().__init__(master)
        self.conexao = conexao
        self.cursor = self.conexao.obter_cursor()
        
        

        self.style = ttk.Style()
        self.style.configure('TLabel', padding=(5, 5), font=('Arial', 10))
        self.style.configure('TEntry', padding=(5, 5), font=('Arial', 10))
        self.style.configure('TButton', padding=(10, 5), font=('Arial', 10))


        # Rótulos e campos de entrada estilizados
        ttk.Label(self, text="Nome:", style="TLabel").grid(row=0, column=1, padx=5, pady=5, sticky=tk.W)
        self.entry_nome = ttk.Entry(self, style="TEntry")
        self.entry_nome.grid(row=0, column=2, padx=5, pady=5)

        ttk.Label(self, text="Caixa:", style="TLabel").grid(row=1, column=1, padx=5, pady=5, sticky=tk.W)
        self.entry_caixa = ttk.Entry(self, style="TEntry")
        self.entry_caixa.grid(row=1, column=2, padx=5, pady=5)

        ttk.Label(self, text="Unidade:", style="TLabel").grid(row=2, column=1, padx=5, pady=5, sticky=tk.W)
        self.entry_unidade = ttk.Entry(self, style="TEntry")
        self.entry_unidade.grid(row=2, column=2, padx=5, pady=5)

        ttk.Label(self, text="Quilo:", style="TLabel").grid(row=3, column=1, padx=5, pady=5, sticky=tk.W)
        self.entry_quilo = ttk.Entry(self, style="TEntry")
        self.entry_quilo.grid(row=3, column=2, padx=5, pady=5)

        ttk.Label(self, text="Litro:", style="TLabel").grid(row=4, column=1, padx=5, pady=5, sticky=tk.W)
        self.entry_litro = ttk.Entry(self, style="TEntry")
        self.entry_litro.grid(row=4, column=2, padx=5, pady=5)

        ttk.Label(self, text="Grama:", style="TLabel").grid(row=5, column=1, padx=5, pady=5, sticky=tk.W)
        self.entry_grama = ttk.Entry(self, style="TEntry")
        self.entry_grama.grid(row=5, column=2, padx=5, pady=5)

        ttk.Label(self, text="Prioridade:", style="TLabel").grid(row=6, column=1, padx=5, pady=5, sticky=tk.W)
        self.entry_prioridade = ttk.Entry(self, style="TEntry")
        self.entry_prioridade.grid(row=6, column=2, padx=5, pady=5)

        # Botão de inserção
        ttk.Button(self, text="Inserir Item", command=self.inserir_item, style="TButton").grid(row=7, column=2, padx=5, pady=5, sticky=tk.W)

    def inserir_item(self):
        # Obtém os valores dos campos de entrada
        nome = self.entry_nome.get()
        caixa = self.entry_caixa.get()
        unidade = self.entry_unidade.get()
        quilo = self.entry_quilo.get()
        litro = self.entry_litro.get()
        grama = self.entry_grama.get()
        prioridade = self.entry_prioridade.get()

        
        # Verifica se a prioridade não é 0
        if int(prioridade) == 0:
            tk.messagebox.showerror("Erro", "A prioridade do item não pode ser 0.")
            return

        # Inserir dados na tabela
        self.cursor.execute('''
            INSERT INTO itens (nome, caixa, unidade, quilo, litro, grama, prioridade)
            VALUES (%s, %s, %s, %s, %s, %s, %s)
        ''', (nome, caixa, unidade, quilo, litro, grama, prioridade))

        # Commit para salvar as alterações no banco de dados
        self.conexao.commit()

        # Exibir mensagem de sucesso
        messagebox.showinfo("Sucesso", "Itens foram inseridos com sucesso.")

        # Limpar os campos de entrada após a inserção
        self.limpar_campos()

    def limpar_campos(self):
        self.entry_nome.delete(0, tk.END)
        self.entry_caixa.delete(0, tk.END)
        self.entry_unidade.delete(0, tk.END)
        self.entry_quilo.delete(0, tk.END)
        self.entry_litro.delete(0, tk.END)
        self.entry_grama.delete(0, tk.END)
        self.entry_prioridade.delete(0, tk.END)

    def fechar_janela(self):
        # Fechar a conexão com o banco de dados antes de fechar a janela
        self.conexao.fechar_conexao()
        self.master.destroy()



class TelaAlteracaoItens(ttk.Frame):
    def __init__(self, master=None, conexao=None):
        super().__init__(master)
        self.conexao = conexao
        self.cursor = self.conexao.obter_cursor()

        # Componentes da interface
        self.style = ttk.Style()
        self.style.configure('TLabel', padding=(5, 5), font=('Arial', 10))
        self.style.configure('TEntry', padding=(5, 5), font=('Arial', 10))
        self.style.configure('TButton', padding=(10, 5), font=('Arial', 10))
        
        self.button_buscar = Button(self, text="Buscar", command=self.buscar_dados)
        self.grid_frame = Frame(self)
        self.grid_frame.grid(row=1, column=1, columnspan=3)



        #Treeview do Estoque para Alteração
        self.tree = ttk.Treeview(self.grid_frame, columns=("Caixa", "Unidade", "Quilo", "Litro", "Grama", "Prioridade"), height=15)
        for column in ("Caixa", "Unidade", "Quilo", "Litro", "Grama", "Prioridade"):
            self.tree.heading(column, text=column)
        self.tree.heading("#0", text="Nome")
        self.tree.heading("Caixa", text="Caixa")
        self.tree.heading("Unidade", text="Unidade")
        self.tree.heading("Quilo", text="Quilo")
        self.tree.heading("Litro", text="Litro")
        self.tree.heading("Grama", text="Grama")
        self.tree.heading("Prioridade", text="Prioridade")


        for column in self.tree["columns"]:
            self.tree.column(column, width=150, anchor='center')


        


        #Componentes existentes
        self.label_nome = ttk.Label(self, text="Nome:", style= "TLabel")
        self.entry_nome = ttk.Entry(self, style= "TEntry")


        self.label_caixa = ttk.Label(self, text="Caixa:", style= "TLabel")
        self.entry_caixa = ttk.Entry(self, style= "TEntry")

        self.label_unidade = ttk.Label(self, text="Unidade:", style= "TLabel")
        self.entry_unidade = ttk.Entry(self, style= "TEntry")

        self.label_quilo = ttk.Label(self, text="Quilo:", style= "TLabel")
        self.entry_quilo = ttk.Entry(self, style= "TEntry")

        self.label_litro = ttk.Label(self, text="Litro:", style= "TLabel")
        self.entry_litro = ttk.Entry(self, style= "TEntry")

        self.label_grama = ttk.Label(self, text="Grama:", style= "TLabel")
        self.entry_grama = ttk.Entry(self, style= "TEntry")

        self.label_prioridade = ttk.Label(self, text="Prioridade:", style= "TLabel")
        self.entry_prioridade = ttk.Entry(self, style= "TEntry")

        self.button_atualizar = ttk.Button(self, text="Atualizar", command=self.atualizar_dados, width= 15)
        self.button_excluir = ttk.Button(self, text="Excluir", command=self.excluir_dados, width= 15)

        self.entry_nome.bind("<Return>", self.exibir_informacoes_completas)
        
        # Posicionamento dos componentes
        self.label_nome.grid(row=2, column=1, sticky="e", padx=(5))
        self.entry_nome.grid(row=2, column=2, pady=5, padx=(0, 5), sticky="w")
        

        # Se necessário, ajuste as outras colunas da Treeview
        self.tree.grid(row=0, column=1, columnspan=3, padx=10, pady=10)
        self.tree.bind("<ButtonRelease-1>", self.selecionar_item)

        self.label_caixa.grid(row=3, column=1, sticky="e", padx=(0, 5), pady=5)
        self.entry_caixa.grid(row=3, column=2, pady=5, sticky= "w")

        self.label_unidade.grid(row=4, column=1, sticky="e", padx=(0, 5), pady=5)
        self.entry_unidade.grid(row=4, column=2, pady=5, sticky= "w")

        self.label_quilo.grid(row=5, column=1, sticky="e", padx=(0, 5), pady=5)
        self.entry_quilo.grid(row=5, column=2, pady=5, sticky= "w")

        self.label_litro.grid(row=6, column=1, sticky="e", padx=(0, 5), pady=5)
        self.entry_litro.grid(row=6, column=2, pady=5, sticky= "w")

        self.label_grama.grid(row=7, column=1, sticky="e", padx=(0, 5), pady=5)
        self.entry_grama.grid(row=7, column=2, pady=5, sticky= "w")

        self.label_prioridade.grid(row=8, column=1, sticky="e", padx=(0, 5), pady=5)
        self.entry_prioridade.grid(row=8, column=2, pady=5, sticky= "w")


        self.button_atualizar.grid(row=9, column=2, pady=5, sticky="w")
        self.button_excluir.grid(row=10, column=2, pady=5, sticky="w")

        # Adicione este atributo à inicialização
        self.atualizar_periodicamente()


        # Configuração do grid
        self.columnconfigure(0, weight=0)
        self.rowconfigure(1, weight=0)

        # Carregar dados iniciais no grid
        self.carregar_dados_iniciais()

    def carregar_dados_iniciais(self):
        # Limpar o Treeview antes de carregar os dados
        for item in self.tree.get_children():
            self.tree.delete(item)

        query = "SELECT Nome, Caixa, Unidade, Quilo, Litro, Grama, Prioridade FROM itens ORDER BY Prioridade ASC"
        self.cursor.execute(query)
        resultados = self.cursor.fetchall()

        for resultado in resultados:
            self.tree.insert("", "end", text=resultado[0], values=resultado[1:])

    def buscar_dados(self):
        nome = self.entry_nome.get()
        query = f"SELECT Caixa, Unidade, Quilo, Litro, Grama, Prioridade FROM itens WHERE Nome = '{nome}'"
        self.cursor.execute(query)
        resultado = self.cursor.fetchone()

        if resultado:
            self.entry_caixa.delete(0, "end")
            self.entry_caixa.insert(0, resultado[0])

            self.entry_unidade.delete(0, "end")
            self.entry_unidade.insert(0, resultado[1])

            self.entry_quilo.delete(0, "end")
            self.entry_quilo.insert(0, resultado[2])

            self.entry_litro.delete(0, "end")
            self.entry_litro.insert(0, resultado[3])

            self.entry_grama.delete(0, "end")
            self.entry_grama.insert(0, resultado[4])

            self.entry_prioridade.delete(0, "end")
            self.entry_prioridade.insert(0, resultado[5])

    def selecionar_item(self, event):
        item_selecionado = self.tree.selection()
        if item_selecionado:
            nome = self.tree.item(item_selecionado, "text")
            self.entry_nome.delete(0, "end")
            self.entry_nome.insert(0, nome)
            self.buscar_dados()

    def atualizar_dados(self):
        nome = self.entry_nome.get()
        caixa = self.entry_caixa.get()
        unidade = self.entry_unidade.get()
        quilo = self.entry_quilo.get()
        litro = self.entry_litro.get()
        grama = self.entry_grama.get()
        prioridade = self.entry_prioridade.get()

        query = f"UPDATE itens SET Caixa='{caixa}', Unidade='{unidade}', Quilo='{quilo}', Litro='{litro}', Grama='{grama}', Prioridade='{prioridade}' WHERE Nome='{nome}'"
        self.cursor.execute(query)
        self.conexao.commit()

    def atualizar_periodicamente(self):
        # Chama carregar_dados_iniciais a cada 1000 milissegundos (1 segundo)
        self.carregar_dados_iniciais()
        self.after(1000, self.atualizar_periodicamente)

    def exibir_informacoes_completas(self, event):
        # Função para exibir informações adicionais com base no nome inserido
        nome = self.entry_nome.get()

        # Verifica se o nome está presente
        if nome:
            query = f"SELECT Caixa, Unidade, Quilo, Litro, Grama, Prioridade FROM itens WHERE Nome = '{nome}'"
            self.cursor.execute(query)
            resultado = self.cursor.fetchone()

            if resultado:
                # Preencher os campos com as informações adicionais
                self.entry_caixa.delete(0, "end")
                self.entry_caixa.insert(0, resultado[0])

                self.entry_unidade.delete(0, "end")
                self.entry_unidade.insert(0, resultado[1])

                self.entry_quilo.delete(0, "end")
                self.entry_quilo.insert(0, resultado[2])

                self.entry_litro.delete(0, "end")
                self.entry_litro.insert(0, resultado[3])

                self.entry_grama.delete(0, "end")
                self.entry_grama.insert(0, resultado[4])

                self.entry_prioridade.delete(0, "end")
                self.entry_prioridade.insert(0, resultado[5])
            else:
                # Limpar os campos se não houver informações
                self.entry_caixa.delete(0, "end")
                self.entry_unidade.delete(0, "end")
                self.entry_quilo.delete(0, "end")
                self.entry_litro.delete(0, "end")
                self.entry_grama.delete(0, "end")
                self.entry_prioridade.delete(0, "end")


    def excluir_dados(self):
        nome = self.entry_nome.get()

        if nome:
            query = f"DELETE FROM itens WHERE Nome='{nome}'"
            self.cursor.execute(query)
            self.conexao.commit()

            # Limpar o Treeview
            for item in self.tree.get_children():
                self.tree.delete(item)

            # Carregar dados atualizados no Treeview
            self.carregar_dados_iniciais()

            # Limpar os campos após a exclusão
            self.entry_nome.delete(0, "end")
            self.entry_caixa.delete(0, "end")
            self.entry_unidade.delete(0, "end")
            self.entry_quilo.delete(0, "end")
            self.entry_litro.delete(0, "end")
            self.entry_grama.delete(0, "end")
            self.entry_prioridade.delete(0, "end")

    def fechar_janela(self):
        # Fechar a conexão com o banco de dados antes de fechar a janela
        self.conexao.conexao.close()
        self.master.destroy()

class TelaInserirLista(tk.Frame):
    def __init__(self, master=None, conexao=None):
        super().__init__(master)
        self.conexao = conexao
        self.cursor = self.conexao.obter_cursor()

        #Estilo dos Componentes
        self.style = ttk.Style()
        self.style.configure('TLabel', padding=(5, 5), font=('Arial', 10))
        self.style.configure('TEntry', padding=(5, 5), font=('Arial', 10))
        self.style.configure('TButton', padding=(10, 5), font=('Arial', 10))
        self.style.configure('TCombobox', padding=(5, 5), font=('Arial', 10))

    # Criar a primeira Treeview para o estoque
        self.treeview_estoque = ttk.Treeview(self, columns=("Caixa", "Unidade", "Quilo", "Litro", "Grama", "Prioridade"),height= 15)
        self.treeview_estoque.heading("#0", text="Item")
        self.treeview_estoque.heading("Caixa", text="Caixa")
        self.treeview_estoque.heading("Unidade", text="Unidade")
        self.treeview_estoque.heading("Quilo", text="Quilo")
        self.treeview_estoque.heading("Litro", text="Litro")
        self.treeview_estoque.heading("Grama", text="Grama")
        self.treeview_estoque.heading("Prioridade", text="Prioridade")
        self.treeview_estoque.grid(row=0, column=0, pady=10, padx=10, sticky=tk.NSEW)

        # Criar a segunda Treeview para a lista para saída
        self.treeview_saida = ttk.Treeview(self, columns=("Caixa", "Unidade", "Quilo", "Litro", "Grama", "Prioridade"))
        self.treeview_saida.heading("#0", text="Item")
        self.treeview_saida.heading("Caixa", text="Caixa")
        self.treeview_saida.heading("Unidade", text="Unidade")
        self.treeview_saida.heading("Quilo", text="Quilo")
        self.treeview_saida.heading("Litro", text="Litro")
        self.treeview_saida.heading("Grama", text="Grama")
        self.treeview_saida.heading("Prioridade", text="Prioridade")
        self.treeview_saida.grid(row=0, column=2, pady=10, padx=10, sticky=tk.NSEW)

        # Adicionar Combobox
        self.criar_combobox()

        # Adicionar elementos para exibir informações
        self.criar_elementos_exibicao()

        # Adicionar botão para enviar itens para o Treeview
        self.botao_enviar = ttk.Button(self, text="Enviar para a Lista", command=self.enviar_para_grid, width=16, style= "TButton")
        self.botao_enviar.grid(row=9, column=1, pady=10, padx=10, sticky=tk.W)

        # Adicionar evento de clique na Treeview do estoque
        self.treeview_estoque.bind("<ButtonRelease-1>", self.mostrar_dados_selecionados)

        # Botão para exportar para o Word
        self.botao_exportar_word = ttk.Button(self, text="Exportar para Word", command=self.exportar_para_word, width=16, style= "TButton")
        self.botao_exportar_word.grid(row=12, column=1, pady=10, padx=10, sticky=tk.W)

        # Adicionar botão para limpar toda a lista de saída
        self.botao_limpar_lista = ttk.Button(self, text="Limpar Lista", command=self.limpar_lista, width=16, style="TButton")
        self.botao_limpar_lista.grid(row=3, column=2, pady=10, padx=10, sticky=tk.W)

        # Adicionar botão para limpar a unidade selecionada da lista de saída
        self.botao_limpar_unidade = ttk.Button(self, text="Limpar Unidade", command=self.limpar_unidade_selecionada, width=16, style="TButton")
        self.botao_limpar_unidade.grid(row=1, column=2, pady=10, padx=10, sticky=tk.W)

        # Ajustar as colunas para ocupar menos espaço
        for column in ("Caixa", "Unidade", "Quilo", "Litro", "Grama", "Prioridade"):
            self.treeview_estoque.column(column, width=30, anchor="center")
            self.treeview_saida.column(column, width=30, anchor="center")

        # Preencher a Treeview do estoque com os itens do banco de dados
        self.preencher_treeview_estoque()

        # Atualizar a treeview de estoque a cada segundo
        self.atualizar_treeview_estoque_periodicamente()

        # Ordenar a treeview do estoque por prioridade
        self.ordenar_por_prioridade(self.treeview_estoque)
        
    

    def criar_combobox(self):
        # Consultar todos os nomes dos itens
        self.cursor.execute("SELECT Nome FROM itens")
        nomes = [nome[0] for nome in self.cursor.fetchall()]

        # Criar a Combobox
        self.combobox = ttk.Combobox(self, values=nomes, width= 17, style='TCombobox')
        self.combobox.set("Escolha um item")  # Texto padrão exibido na Combobox
        self.combobox.grid(row=1, column=1, pady=10, padx=10, sticky=tk.W)


        # Adicionar evento para a mudança na seleção
        self.combobox.bind("<<ComboboxSelected>>", self.atualizar_informacoes)

        # Adicionar evento para autocompletar enquanto digita
        self.combobox.bind("<KeyPress>", self.autocompletar_combobox)

    def autocompletar_combobox(self, event):
        # Obter o texto digitado na Combobox
        entrada = self.combobox.get().lower()

        # Consultar todos os nomes dos itens
        self.cursor.execute("SELECT Nome FROM itens")
        nomes = [nome[0] for nome in self.cursor.fetchall()]

        # Filtrar nomes que começam com o texto digitado
        sugestoes = [nome for nome in nomes if nome.lower().startswith(entrada)]

        # Atualizar a lista de sugestões da Combobox
        self.combobox['values'] = sugestoes

    def criar_elementos_exibicao(self):
        # Criar Labels e Entries para exibir informações
        self.label_caixa = ttk.Label(self, text="Caixa:", style= "TLabel")
        self.entry_caixa = ttk.Entry(self, style= "TEntry")
        self.label_unidade = ttk.Label(self, text="Unidade:", style= "TLabel")
        self.entry_unidade = ttk.Entry(self, style= "TEntry")
        self.label_quilo = ttk.Label(self, text="Quilo:", style= "TLabel")
        self.entry_quilo = ttk.Entry(self, style= "TEntry")
        self.label_litro = ttk.Label(self, text="Litro:", style= "TLabel")
        self.entry_litro = ttk.Entry(self, style= "TEntry")
        self.label_grama = ttk.Label(self, text="Grama:", style= "TLabel")
        self.entry_grama = ttk.Entry(self, style= "TEntry")
        self.label_prioridade = ttk.Label(self, text="Prioridade:", style= "TLabel")
        self.entry_prioridade = ttk.Entry(self, style= "TEntry")

        # Posicionar Labels e Entries
        self.label_caixa.grid(row=3, column=0, pady=5, padx=10, sticky=tk.E)
        self.entry_caixa.grid(row=3, column=1, pady=5, padx=10, sticky=tk.E)
        self.label_unidade.grid(row=4, column=0, pady=5, padx=10, sticky=tk.E)
        self.entry_unidade.grid(row=4, column=1, pady=5, padx=10, sticky=tk.E)
        self.label_quilo.grid(row=5, column=0, pady=5, padx=10, sticky=tk.E)
        self.entry_quilo.grid(row=5, column=1, pady=5, padx=10, sticky=tk.E)
        self.label_litro.grid(row=6, column=0, pady=5, padx=10, sticky=tk.E)
        self.entry_litro.grid(row=6, column=1, pady=5, padx=10, sticky=tk.E)
        self.label_grama.grid(row=7, column=0, pady=5, padx=10, sticky=tk.E)
        self.entry_grama.grid(row=7, column=1, pady=5, padx=10, sticky=tk.E)
        self.label_prioridade.grid(row=8, column=0, pady=5, padx=10, sticky=tk.E)
        self.entry_prioridade.grid(row=8, column=1, pady=5, padx=10, sticky=tk.E)


    def atualizar_informacoes(self, event):
        # Obter o nome selecionado na ComboBox
        nome_selecionado = self.combobox.get()

        # Consultar informações do item selecionado
        self.cursor.execute("SELECT Caixa, Unidade, Quilo, Litro, Grama, Prioridade FROM itens WHERE Nome=%s", (nome_selecionado,))
        informacoes = self.cursor.fetchone()

        # Preencher as Entries com as informações obtidas
        self.entry_caixa.delete(0, tk.END)
        self.entry_caixa.insert(0, informacoes[0])

        self.entry_unidade.delete(0, tk.END)
        self.entry_unidade.insert(0, informacoes[1])

        self.entry_quilo.delete(0, tk.END)
        self.entry_quilo.insert(0, informacoes[2])

        self.entry_litro.delete(0, tk.END)
        self.entry_litro.insert(0, informacoes[3])

        self.entry_grama.delete(0, tk.END)
        self.entry_grama.insert(0, informacoes[4])

        self.entry_prioridade.delete(0, tk.END)
        self.entry_prioridade.insert(0, informacoes[5])

    def enviar_para_grid(self):
        nome_selecionado = self.combobox.get()

        # Obter os valores alterados das Entries
        caixa = self.entry_caixa.get()
        unidade = self.entry_unidade.get()
        quilo = self.entry_quilo.get()
        litro = self.entry_litro.get()
        grama = self.entry_grama.get()
        prioridade = self.entry_prioridade.get()

        # Validar se a prioridade é um número válido e não zero
        if prioridade and prioridade.isdigit() and int(prioridade) != 0:
            # Preencher a Treeview da lista para saída com os novos valores e organizar por prioridade
            self.treeview_saida.insert("", "end", text=nome_selecionado, values=(caixa, unidade, quilo, litro, grama, prioridade))
            self.ordenar_por_prioridade(self.treeview_saida)
        else:
            messagebox.showwarning("Aviso", "A prioridade deve ser um número válido e diferente de zero.")

    def ordenar_por_prioridade(self, treeview):
        # Organizar o Treeview pela coluna "Prioridade"
        data = [(treeview.item(item)["values"], item) for item in treeview.get_children()]
        data.sort(key=lambda x: int(x[0][5]))  # Ordenar com base na prioridade (índice 5)
        for index, (values, item) in enumerate(data):
            treeview.move(item, "", index)

    def exportar_para_word(self):
        file_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Documentos Word", "*.docx")])
        if file_path:
            # Criar um documento Word
            document = docx.Document()

            # Adicionar cabeçalho para a lista para saída
            document.add_heading('Lista para Saída', level=1)

            # Adicionar dados do Treeview da lista para saída ao documento Word
            self.adicionar_dados_ao_documento(document, self.treeview_saida)

            # Salvar o documento Word no local escolhido pelo usuário
            document.save(file_path)
            messagebox.showinfo("Exportação Concluída", f"Os dados foram exportados para o arquivo {file_path}.")

    def atualizar_treeview_estoque_periodicamente(self):
        # Salvar a ordem atual da Treeview do estoque
        current_order = self.obter_ordem_treeview(self.treeview_estoque)

        # Verificar se a treeview ainda existe
        if self.treeview_estoque.winfo_exists():
            # Atualizar a Treeview do estoque com os dados mais recentes do banco de dados
            self.atualizar_treeview_estoque()

            # Restaurar a ordem original da Treeview do estoque
            self.restaurar_ordem_treeview(self.treeview_estoque, current_order)

        # Agendar a próxima atualização após 1000 milissegundos (1 segundo)
        self.after(1000, self.atualizar_treeview_estoque_periodicamente)

    def obter_ordem_treeview(self, treeview):
        # Obter a ordem atual dos itens na Treeview
        return [item for item in treeview.get_children() if treeview.exists(item)]

    def restaurar_ordem_treeview(self, treeview, order):
        # Restaurar a ordem original dos itens na Treeview
        for item in order:
            if treeview.exists(item):
                treeview.move(item, "", "end")

    def atualizar_treeview_estoque(self):
        # Limpar a Treeview do estoque
        for item in self.treeview_estoque.get_children():
            self.treeview_estoque.delete(item)

        # Consultar todos os itens do banco de dados, ordenando por prioridade
        self.cursor.execute("SELECT Nome, Caixa, Unidade, Quilo, Litro, Grama, Prioridade FROM itens ORDER BY Prioridade ASC")
        for row in self.cursor.fetchall():
            nome, caixa, unidade, quilo, litro, grama, prioridade = row

            # Verificar se a unidade é diferente de '0' antes de adicionar à Treeview
            if unidade != '0':
                self.treeview_estoque.insert("", "end", text=nome, values=(caixa, unidade, quilo, litro, grama, prioridade))

    def limpar_lista(self):
        # Limpar toda a Treeview da lista para saída
        for item in self.treeview_saida.get_children():
            self.treeview_saida.delete(item)

    def limpar_unidade_selecionada(self):
        # Obter o item selecionado na Treeview de saída
        item_selecionado = self.treeview_saida.selection()

        # Verificar se algum item foi selecionado
        if item_selecionado:
            # Remover o item selecionado na Treeview de saída
            self.treeview_saida.delete(item_selecionado)

    def adicionar_dados_ao_documento(self, document, treeview):
        for item in treeview.get_children():
            nome = treeview.item(item, "text")
            values = treeview.item(item, "values")

            # Verificar se caixas, unidades, quilos, litros e gramas são diferentes de "0"
            detalhes_itens = []

            if values[0] != "0":
                detalhes_itens.append(f'{values[0]} Caixas')
            if values[1] != "0":
                detalhes_itens.append(f'{values[1]} Unidades')
            if values[2] != "0":
                detalhes_itens.append(f'{values[2]} Quilos')
            if values[3] != "0":
                detalhes_itens.append(f'{values[3]} Litros')
            if values[4] != "0":
                detalhes_itens.append(f'{values[4]} Gramas')

            # Construir a string de detalhes apenas se houver pelo menos um valor não zero
            detalhes = ', '.join(detalhes_itens)
            
            if detalhes:  # Adiciona uma quebra de linha apenas se houver detalhes
                detalhes = f'{nome}: {detalhes}'
                document.add_paragraph(detalhes)

                    


    def mostrar_dados_selecionados(self, event):
        # Obter o item selecionado na Treeview do estoque
        item_selecionado = self.treeview_estoque.selection()

        if item_selecionado:
            # Obter o nome do item selecionado
            nome_selecionado = self.treeview_estoque.item(item_selecionado, "text")

            # Definir o nome selecionado na Combobox
            self.combobox.set(nome_selecionado)

            # Consultar informações do item selecionado
            self.cursor.execute("SELECT Caixa, Unidade, Quilo, Litro, Grama, Prioridade FROM itens WHERE Nome=%s", (nome_selecionado,))
            informacoes = self.cursor.fetchone()

            # Preencher as Entries com as informações obtidas
            self.entry_caixa.delete(0, tk.END)
            self.entry_caixa.insert(0, informacoes[0])

            self.entry_unidade.delete(0, tk.END)
            self.entry_unidade.insert(0, informacoes[1])

            self.entry_quilo.delete(0, tk.END)
            self.entry_quilo.insert(0, informacoes[2])

            self.entry_litro.delete(0, tk.END)
            self.entry_litro.insert(0, informacoes[3])

            self.entry_grama.delete(0, tk.END)
            self.entry_grama.insert(0, informacoes[4])

            self.entry_prioridade.delete(0, tk.END)
            self.entry_prioridade.insert(0, informacoes[5])

    def preencher_treeview_estoque(self):
        # Limpar a Treeview do estoque
        for item in self.treeview_estoque.get_children():
            self.treeview_estoque.delete(item)

        # Consultar todos os itens do banco de dados
        self.cursor.execute("SELECT Nome, Caixa, Unidade, Quilo, Litro, Grama, Prioridade FROM itens")
        for row in self.cursor.fetchall():
            nome, caixa, unidade, quilo, litro, grama, prioridade = row

            # Verificar se a unidade é diferente de '0' antes de adicionar à Treeview
            if unidade != '0':
                self.treeview_estoque.insert("", "end", text=nome, values=(caixa, unidade, quilo, litro, grama, prioridade))

class TelaInsercaoAlimentos(ttk.Frame):
    def __init__(self, master=None, conexao=None):
        super().__init__(master)
        self.conexao = conexao
        self.cursor = self.conexao.obter_cursor()

        self.style = ttk.Style()
        self.style.configure('TLabel', padding=(5, 5), font=('Arial', 10))
        self.style.configure('TEntry', padding=(5, 5), font=('Arial', 10))
        self.style.configure('TButton', padding=(10, 5), font=('Arial', 10))

        # Rótulos e campos de entrada estilizados
        ttk.Label(self, text="Nome:", style="TLabel").grid(row=0, column=1, padx=5, pady=5, sticky=tk.W)
        self.entry_nome = ttk.Entry(self, style="TEntry")
        self.entry_nome.grid(row=0, column=2, padx=5, pady=5)

        ttk.Label(self, text="Caixa:", style="TLabel").grid(row=1, column=1, padx=5, pady=5, sticky=tk.W)
        self.entry_caixa = ttk.Entry(self, style="TEntry")
        self.entry_caixa.grid(row=1, column=2, padx=5, pady=5)

        ttk.Label(self, text="Unidade:", style="TLabel").grid(row=2, column=1, padx=5, pady=5, sticky=tk.W)
        self.entry_unidade = ttk.Entry(self, style="TEntry")
        self.entry_unidade.grid(row=2, column=2, padx=5, pady=5)

        ttk.Label(self, text="Quilo:", style="TLabel").grid(row=3, column=1, padx=5, pady=5, sticky=tk.W)
        self.entry_quilo = ttk.Entry(self, style="TEntry")
        self.entry_quilo.grid(row=3, column=2, padx=5, pady=5)

        ttk.Label(self, text="Litro:", style="TLabel").grid(row=4, column=1, padx=5, pady=5, sticky=tk.W)
        self.entry_litro = ttk.Entry(self, style="TEntry")
        self.entry_litro.grid(row=4, column=2, padx=5, pady=5)

        ttk.Label(self, text="Grama:", style="TLabel").grid(row=5, column=1, padx=5, pady=5, sticky=tk.W)
        self.entry_grama = ttk.Entry(self, style="TEntry")
        self.entry_grama.grid(row=5, column=2, padx=5, pady=5)

        ttk.Label(self, text="Prioridade:", style="TLabel").grid(row=6, column=1, padx=5, pady=5, sticky=tk.W)
        self.entry_prioridade = ttk.Entry(self, style="TEntry")
        self.entry_prioridade.grid(row=6, column=2, padx=5, pady=5)

        # Botão de inserção
        ttk.Button(self, text="Inserir Alimento", command=self.inserir_alimento, style="TButton").grid(row=7, column=2, padx=5, pady=5, sticky=tk.W)

    def inserir_alimento(self):
        nome = self.entry_nome.get()
        caixa = self.entry_caixa.get()
        unidade = self.entry_unidade.get()
        quilo = self.entry_quilo.get()
        litro = self.entry_litro.get()
        grama = self.entry_grama.get()
        prioridade = self.entry_prioridade.get()

        if int(prioridade) == 0:
            tk.messagebox.showerror("Erro", "A prioridade do alimento não pode ser 0.")
            return

        # Inserir dados na tabela
        self.cursor.execute('''
            INSERT INTO alimentos (Nome, Caixa, Unidade, Quilo, Litro, Grama, Prioridade)
            VALUES (%s, %s, %s, %s, %s, %s, %s)
        ''', (nome, caixa, unidade, quilo, litro, grama, prioridade))

        # Commit para salvar as alterações no banco de dados
        self.conexao.commit()

        messagebox.showinfo("Sucesso", "Alimento inserido com sucesso.")
        self.limpar_campos()

    def limpar_campos(self):
        self.entry_nome.delete(0, tk.END)
        self.entry_caixa.delete(0, tk.END)
        self.entry_unidade.delete(0, tk.END)
        self.entry_quilo.delete(0, tk.END)
        self.entry_litro.delete(0, tk.END)
        self.entry_grama.delete(0, tk.END)
        self.entry_prioridade.delete(0, tk.END)

    def fechar_janela(self):
        self.conexao.fechar_conexao()
        self.master.destroy()


class TelaAlteracaoAlimentos(ttk.Frame):
    def __init__(self, master=None, conexao=None):
        super().__init__(master)
        self.conexao = conexao
        self.cursor = self.conexao.obter_cursor()

        # Componentes da interface
        self.style = ttk.Style()
        self.style.configure('TLabel', padding=(5, 5), font=('Arial', 10))
        self.style.configure('TEntry', padding=(5, 5), font=('Arial', 10))
        self.style.configure('TButton', padding=(10, 5), font=('Arial', 10))

        self.grid_frame = Frame(self)
        self.grid_frame.grid(row=1, column=1, columnspan=3)

        # Treeview do Estoque para Alteração
        self.tree = ttk.Treeview(self.grid_frame, columns=("Caixa", "Unidade", "Quilo", "Litro", "Grama", "Prioridade"), height=15)
        for column in ("Caixa", "Unidade", "Quilo", "Litro", "Grama", "Prioridade"):
            self.tree.heading(column, text=column)
        self.tree.heading("#0", text="Nome")
        self.tree.heading("Caixa", text="Caixa")
        self.tree.heading("Unidade", text="Unidade")
        self.tree.heading("Quilo", text="Quilo")
        self.tree.heading("Litro", text="Litro")
        self.tree.heading("Grama", text="Grama")
        self.tree.heading("Prioridade", text="Prioridade")

        for column in self.tree["columns"]:
            self.tree.column(column, width=150, anchor='center')

        # Componentes existentes
        self.label_nome = ttk.Label(self, text="Nome:", style="TLabel")
        self.entry_nome = ttk.Entry(self, style="TEntry")

        self.label_caixa = ttk.Label(self, text="Caixa:", style="TLabel")
        self.entry_caixa = ttk.Entry(self, style="TEntry")

        self.label_unidade = ttk.Label(self, text="Unidade:", style="TLabel")
        self.entry_unidade = ttk.Entry(self, style="TEntry")

        self.label_quilo = ttk.Label(self, text="Quilo:", style="TLabel")
        self.entry_quilo = ttk.Entry(self, style="TEntry")

        self.label_litro = ttk.Label(self, text="Litro:", style="TLabel")
        self.entry_litro = ttk.Entry(self, style="TEntry")

        self.label_grama = ttk.Label(self, text="Grama:", style="TLabel")
        self.entry_grama = ttk.Entry(self, style="TEntry")

        self.label_prioridade = ttk.Label(self, text="Prioridade:", style="TLabel")
        self.entry_prioridade = ttk.Entry(self, style="TEntry")

        self.button_atualizar = ttk.Button(self, text="Atualizar", command=self.atualizar_dados, width=15)
        self.button_excluir = ttk.Button(self, text="Excluir", command=self.excluir_dados, width=15)

        self.entry_nome.bind("<Return>", self.exibir_informacoes_completas)

        # Posicionamento dos componentes
        self.label_nome.grid(row=2, column=1, sticky="e", padx=(5))
        self.entry_nome.grid(row=2, column=2, pady=5, padx=(0, 5), sticky="w")

        self.tree.grid(row=0, column=1, columnspan=3, padx=10, pady=10)
        self.tree.bind("<ButtonRelease-1>", self.selecionar_item)

        self.label_caixa.grid(row=3, column=1, sticky="e", padx=(0, 5), pady=5)
        self.entry_caixa.grid(row=3, column=2, pady=5, sticky="w")

        self.label_unidade.grid(row=4, column=1, sticky="e", padx=(0, 5), pady=5)
        self.entry_unidade.grid(row=4, column=2, pady=5, sticky="w")

        self.label_quilo.grid(row=5, column=1, sticky="e", padx=(0, 5), pady=5)
        self.entry_quilo.grid(row=5, column=2, pady=5, sticky="w")

        self.label_litro.grid(row=6, column=1, sticky="e", padx=(0, 5), pady=5)
        self.entry_litro.grid(row=6, column=2, pady=5, sticky="w")

        self.label_grama.grid(row=7, column=1, sticky="e", padx=(0, 5), pady=5)
        self.entry_grama.grid(row=7, column=2, pady=5, sticky="w")

        self.label_prioridade.grid(row=8, column=1, sticky="e", padx=(0, 5), pady=5)
        self.entry_prioridade.grid(row=8, column=2, pady=5, sticky="w")

        self.button_atualizar.grid(row=9, column=2, pady=5, sticky="w")
        self.button_excluir.grid(row=10, column=2, pady=5, sticky="w")

        # Adicione este atributo à inicialização
        self.atualizar_periodicamente()

        # Configuração do grid
        self.columnconfigure(0, weight=0)
        self.rowconfigure(1, weight=0)

        # Carregar dados iniciais no grid
        self.carregar_dados_iniciais()

    def carregar_dados_iniciais(self):
        for item in self.tree.get_children():
            self.tree.delete(item)

        query = "SELECT Nome, Caixa, Unidade, Quilo, Litro, Grama, Prioridade FROM alimentos ORDER BY Prioridade ASC"
        self.cursor.execute(query)
        resultados = self.cursor.fetchall()

        for resultado in resultados:
            self.tree.insert("", "end", text=resultado[0], values=resultado[1:])

    def buscar_dados(self):
        nome = self.entry_nome.get()
        query = f"SELECT Caixa, Unidade, Quilo, Litro, Grama, Prioridade FROM alimentos WHERE Nome = '{nome}'"
        self.cursor.execute(query)
        resultado = self.cursor.fetchone()

        if resultado:
            self.entry_caixa.delete(0, "end")
            self.entry_caixa.insert(0, resultado[0])

            self.entry_unidade.delete(0, "end")
            self.entry_unidade.insert(0, resultado[1])

            self.entry_quilo.delete(0, "end")
            self.entry_quilo.insert(0, resultado[2])

            self.entry_litro.delete(0, "end")
            self.entry_litro.insert(0, resultado[3])

            self.entry_grama.delete(0, "end")
            self.entry_grama.insert(0, resultado[4])

            self.entry_prioridade.delete(0, "end")
            self.entry_prioridade.insert(0, resultado[5])

    def selecionar_item(self, event):
        item_selecionado = self.tree.selection()
        if item_selecionado:
            nome = self.tree.item(item_selecionado, "text")
            self.entry_nome.delete(0, "end")
            self.entry_nome.insert(0, nome)
            self.buscar_dados()

    def atualizar_dados(self):
        nome = self.entry_nome.get()
        caixa = self.entry_caixa.get()
        unidade = self.entry_unidade.get()
        quilo = self.entry_quilo.get()
        litro = self.entry_litro.get()
        grama = self.entry_grama.get()
        prioridade = self.entry_prioridade.get()

        query = f"UPDATE alimentos SET Caixa='{caixa}', Unidade='{unidade}', Quilo='{quilo}', Litro='{litro}', Grama='{grama}', Prioridade='{prioridade}' WHERE Nome='{nome}'"
        self.cursor.execute(query)
        self.conexao.commit()

    def atualizar_periodicamente(self):
        self.carregar_dados_iniciais()
        self.after(1000, self.atualizar_periodicamente)

    def exibir_informacoes_completas(self, event):
        nome = self.entry_nome.get()

        if nome:
            query = f"SELECT Caixa, Unidade, Quilo, Litro, Grama, Prioridade FROM alimentos WHERE Nome = '{nome}'"
            self.cursor.execute(query)
            resultado = self.cursor.fetchone()

            if resultado:
                self.entry_caixa.delete(0, "end")
                self.entry_caixa.insert(0, resultado[0])

                self.entry_unidade.delete(0, "end")
                self.entry_unidade.insert(0, resultado[1])

                self.entry_quilo.delete(0, "end")
                self.entry_quilo.insert(0, resultado[2])

                self.entry_litro.delete(0, "end")
                self.entry_litro.insert(0, resultado[3])

                self.entry_grama.delete(0, "end")
                self.entry_grama.insert(0, resultado[4])

                self.entry_prioridade.delete(0, "end")
                self.entry_prioridade.insert(0, resultado[5])
            else:
                self.entry_caixa.delete(0, "end")
                self.entry_unidade.delete(0, "end")
                self.entry_quilo.delete(0, "end")
                self.entry_litro.delete(0, "end")
                self.entry_grama.delete(0, "end")
                self.entry_prioridade.delete(0, "end")

    def excluir_dados(self):
        nome = self.entry_nome.get()

        if nome:
            query = f"DELETE FROM alimentos WHERE Nome='{nome}'"
            self.cursor.execute(query)
            self.conexao.commit()

            for item in self.tree.get_children():
                self.tree.delete(item)

            self.carregar_dados_iniciais()

            self.entry_nome.delete(0, "end")
            self.entry_caixa.delete(0, "end")
            self.entry_unidade.delete(0, "end")
            self.entry_quilo.delete(0, "end")
            self.entry_litro.delete(0, "end")
            self.entry_grama.delete(0, "end")
            self.entry_prioridade.delete(0, "end")

    def fechar_janela(self):
        self.conexao.conexao.close()
        self.master.destroy()


class AplicacaoPrincipal:
    def __init__(self, root):
        self.root = root
        self.root.title("Sistema de Estoque")

        self.root.tk.call("source", theme_path)
        root.tk.call("set_theme", "light")
        
        
        # Configuração da conexão com o banco de dados MySQL
        self.conexao = ConexaoBancoDados()

        # Criar o widget Notebook para as abas
        self.notebook = ttk.Notebook(root)
        self.notebook.pack(fill='both', expand=True)

        # Criar as abas e adicionar ao Notebook
        self.aba_insercao = TelaInsercaoItens(self.notebook, self.conexao)
        self.notebook.add(self.aba_insercao, text="Inserção de Itens")

        self.aba_alteracao = TelaAlteracaoItens(self.notebook, self.conexao)
        self.notebook.add(self.aba_alteracao, text="Alteração de Itens")

        self.aba_lista = TelaInserirLista(self.notebook, self.conexao)
        self.notebook.add(self.aba_lista, text="Lista - Geral")

        self.aba_insercao = TelaInsercaoAlimentos(self.notebook, self.conexao)
        self.notebook.add(self.aba_insercao, text="Inserção de Alimentos")

        self.aba_insercao = TelaAlteracaoAlimentos(self.notebook, self.conexao)
        self.notebook.add(self.aba_insercao, text="Alteração de Alimentos")

        root.iconbitmap(default=icon_path)

    def posicionar_janela(self):  
        # Posicionar a janela no topo e mais para o lado direito, ajuste conforme necessário
        altura_tela_inical = 40
        distancia_lado_inicial = 400
        self.root.geometry(f"+{distancia_lado_inicial}+{altura_tela_inical}")

if __name__ == "__main__":
    root = tk.Tk()
    app = AplicacaoPrincipal(root)
    app.posicionar_janela()
    root.resizable(width=False, height=False)
    root.mainloop()
